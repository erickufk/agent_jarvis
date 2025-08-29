"""tools_local.py
Local tool implementations.
Author: Uriel Kaiia (kaya.uf@gmail.com)


All custom tools and sandbox utilities (e.g., file I/O, OS helpers)
are defined here, following the MCP-style separation of concerns.

Design:
- ACTION_ROOT: global sandbox folder the agent can touch
- LAST_EXPLORER_PICK: last item clicked in FileExplorer (file or folder)
- get_tools(safe_mode): returns list of @tool functions (no shell in safe mode)
"""

from pathlib import Path
from typing import Optional, Iterable
import os, re
from collections import Counter

from PIL import Image, ImageOps, ImageFilter
import pytesseract
from smolagents import tool

# --------- Global sandbox state ----------
ACTION_ROOT: Path = Path.cwd().resolve()
LAST_EXPLORER_PICK: Optional[Path] = None

# OCR / PDF helpers (env-configurable in app startup)
DEFAULT_OCR_LANGS = os.getenv("OCR_LANGS", "eng")
POPPLER_PATH = os.getenv("POPPLER_PATH")  # for pdf2image on Windows
# Default Tesseract path (Windows). If set, wire it; otherwise let pytesseract find it.
_TESS_EXE = os.getenv("TESSERACT_CMD", r"C:\Program Files\Tesseract-OCR\tesseract.exe")
if os.path.exists(_TESS_EXE):
    pytesseract.pytesseract.tesseract_cmd = _TESS_EXE

# --------- Folder setters / helpers ----------
def _extract_path_from_explorer_selection(selected) -> Optional[Path]:
    """Robustly parse FileExplorer selection across Gradio versions."""
    if not selected:
        return None

    def _one(x):
        if x is None:
            return None
        if isinstance(x, (str, Path)):
            p = Path(x)
            return p if p.is_absolute() else (Path.cwd() / p)
        if isinstance(x, dict):
            # Common shapes: {"path": "..."} or {"name": "..."}
            if x.get("path"):
                p = Path(x["path"])
                return p if p.is_absolute() else (Path.cwd() / p)
            if x.get("name"):
                return (Path.cwd() / str(x["name"]))
        return None

    if isinstance(selected, list) and selected:
        return _one(selected[0])
    return _one(selected)

def set_action_root_from_any(selected) -> str:
    """Set ACTION_ROOT from FileExplorer selection (file → parent)."""
    global ACTION_ROOT, LAST_EXPLORER_PICK
    p = _extract_path_from_explorer_selection(selected)
    if p is None:
        return str(ACTION_ROOT)
    LAST_EXPLORER_PICK = p
    ACTION_ROOT = p if p.is_dir() else p.parent
    return str(ACTION_ROOT)

def set_action_root_from_text(path_str: str) -> str:
    """Set ACTION_ROOT from a textbox path."""
    global ACTION_ROOT
    p = Path(path_str).expanduser().resolve()
    if not p.exists():
        raise FileNotFoundError(f"path not found: {p}")
    ACTION_ROOT = p if p.is_dir() else p.parent
    return str(ACTION_ROOT)

def get_action_root() -> str:
    """Return ACTION_ROOT as a string (for UI)."""
    return str(ACTION_ROOT)

def _normalize_to_root(user_path: str) -> Path:
    """Resolve a path (absolute or relative) under ACTION_ROOT; block escapes."""
    p = Path(user_path)
    rp = (ACTION_ROOT / p).resolve() if not p.is_absolute() else p.resolve()
    if ACTION_ROOT == rp or ACTION_ROOT in rp.parents:
        return rp
    raise ValueError(f"Path escapes sandbox: {rp} not under {ACTION_ROOT}")

# ---- dedup helpers for recursive scans (avoid symlink/junction double counts)
def _excluded(rel_posix: str, exclude_names: set[str]) -> bool:
    return any(f"/{x}/" in f"/{rel_posix}/" or rel_posix.startswith(f"{x}/") for x in exclude_names)

def _iter_files_unique(base: Path, pattern: str, exclude: Iterable[str]) -> Iterable[Path]:
    """
    Yield files under `base` matching `pattern`, excluding listed dir names,
    de-duplicated by realpath (prevents double-counting via junctions/symlinks).
    """
    exclude_names = {x.strip() for x in exclude if x and x.strip()}
    seen_real = set()
    it = base.rglob(pattern) if "**" in pattern else base.glob(pattern)
    for p in it:
        try:
            if not p.is_file():
                continue
            rel = p.relative_to(ACTION_ROOT).as_posix()
            if _excluded(rel, exclude_names):
                continue
            real = p.resolve(strict=False)
            key = str(real).lower()  # lower() to avoid Win case dupes
            if key in seen_real:
                continue
            seen_real.add(key)
            yield p
        except Exception:
            continue

# --------- General tools ----------
@tool
def cwd() -> str:
    """Return the absolute path of the current action folder.

    Returns:
        str: Absolute path to ACTION_ROOT.
    """
    return str(ACTION_ROOT)

@tool
def write_file(path: str, content: str) -> str:
    """Create or overwrite a UTF-8 text file (sandboxed to ACTION_ROOT).

    Args:
        path (str): Relative path inside the action folder where the file is written.
        content (str): Text content to write.

    Returns:
        str: 'saved:<relative-path>' on success or 'error:<message>' on failure.
    """
    try:
        target = _normalize_to_root(path)
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(content, encoding="utf-8")
        return f"saved:{target.relative_to(ACTION_ROOT)}"
    except Exception as e:
        return f"error:{e}"

@tool
def read_text(path: str, max_chars: int = 20000) -> str:
    """Read a text file from the action folder with robust decoding.

    Args:
        path (str): Relative/absolute file path under action folder.
        max_chars (int): Max characters to return.

    Returns:
        str: File contents (possibly truncated) or 'error:<message>'.
    """
    try:
        target = _normalize_to_root(path)
        data = target.read_bytes()
        for enc in ("utf-8", "cp1251", "cp866", "utf-16", "latin-1"):
            try:
                return data.decode(enc)[:max_chars]
            except Exception:
                continue
        return data.decode("utf-8", errors="replace")[:max_chars]
    except Exception as e:
        return f"error:{e}"

@tool
def list_dir(path: str = ".", glob_pattern: str = "*", exclude: str = "") -> str:
    """List files/folders that match a glob pattern.

    Args:
        path (str): Subfolder under the action folder.
        glob_pattern (str): Glob to match (e.g., '*.md', '**/*.py').
        exclude (str): Comma-separated dir names to skip anywhere in the path.

    Returns:
        str: Matches relative to action folder, or '(no matches)'.
    """
    try:
        base = _normalize_to_root(path)
        exclude_names = {x.strip() for x in exclude.split(",") if x.strip()}
        # files (deduped)
        files = [
            str(p.relative_to(ACTION_ROOT))
            for p in _iter_files_unique(base, glob_pattern, exclude_names)
        ]
        # dirs (non-deduped, just to show structure)
        dirs = []
        it = base.rglob(glob_pattern) if "**" in glob_pattern else base.glob(glob_pattern)
        for p in it:
            if p.is_dir():
                rel = p.relative_to(ACTION_ROOT).as_posix()
                if not _excluded(rel, exclude_names):
                    dirs.append(rel)
        out = sorted(set(files)) + sorted(set(dirs))
        return "\n".join(out) if out else "(no matches)"
    except Exception as e:
        return f"error:{e}"

@tool
def search_text(path: str = ".", glob_pattern: str = "**/*", query: str = "", max_hits: int = 100) -> str:
    """Case-insensitive regex search within files.

    Args:
        path (str): Subfolder under the action folder.
        glob_pattern (str): File glob to include (e.g., '**/*.py').
        query (str): Regular expression (case-insensitive). Must be non-empty.
        max_hits (int): Max number of matched lines.

    Returns:
        str: '<relative-path>:<line-no>: <line>' lines or '(no hits)' / 'error:<msg>'.
    """
    if not query or not query.strip():
        return "error: query must be a non-empty regex"
    try:
        rx = re.compile(query, re.IGNORECASE)
        root = _normalize_to_root(path)
        hits = []
        for f in root.rglob(glob_pattern):
            if not f.is_file():
                continue
            try:
                for i, line in enumerate(f.read_text(encoding="utf-8", errors="ignore").splitlines(), 1):
                    if rx.search(line):
                        hits.append(f"{f.relative_to(ACTION_ROOT)}:{i}: {line.strip()}")
                        if len(hits) >= max_hits:
                            return "\n".join(hits)
            except Exception:
                continue
        return "\n".join(hits) if hits else "(no hits)"
    except Exception as e:
        return f"error:{e}"

@tool
def sh(cmd: str, shell: str = "auto") -> str:
    """Execute a shell command inside the action folder with UTF-8 output.

    Args:
        cmd (str): Command to execute (Windows: 'dir', 'type', 'findstr').
        shell (str): 'auto'|'cmd'|'powershell' (Windows only).

    Returns:
        str: stdout+stderr (UTF-8) or 'error:<message>'.
    """
    try:
        import subprocess, os as _os
        cmd_lower = cmd.strip().lower()
        if cmd_lower in ("pwd", "cd", "cwd"):
            return str(ACTION_ROOT)
        if cmd_lower in ("printenv", "env"):
            keys = ["USERNAME", "USERPROFILE", "HOMEDRIVE", "HOMEPATH", "USERDOMAIN", "COMPUTERNAME", "PATH"]
            return "\n".join(f"{k}={_os.environ.get(k, '')}" for k in keys)

        if _os.name == "nt":
            if shell in ("auto", "cmd"):
                cmd_win = f"chcp 65001 >NUL & {cmd}"
                r = subprocess.run(
                    cmd_win, shell=True, cwd=str(ACTION_ROOT),
                    capture_output=True, text=True, encoding="utf-8", errors="replace"
                )
                return (r.stdout or "") + (r.stderr or "")
            else:
                ps_cmd = f"[Console]::OutputEncoding=[System.Text.UTF8Encoding]::new(); {cmd}"
                r = subprocess.run(
                    ["powershell", "-NoProfile", "-Command", ps_cmd],
                    cwd=str(ACTION_ROOT), capture_output=True, text=True, encoding="utf-8", errors="replace"
                )
                return (r.stdout or "") + (r.stderr or "")
        else:
            r = subprocess.run(
                cmd, shell=True, cwd=str(ACTION_ROOT),
                capture_output=True, text=True, encoding="utf-8", errors="replace"
            )
            return (r.stdout or "") + (r.stderr or "")
    except Exception as e:
        return f"error:{e}"

# --------- Document formats ----------
@tool
def read_pdf_auto(
    path: str,
    max_pages: int = 10,
    dpi: int = 300,
    lang: str = "",
    preprocess: str = "auto",
    max_chars: int = 20000,
) -> str:
    """Extract text from a PDF with an automatic fallback to OCR.

    Args:
        path (str): PDF path under the action folder (relative to ACTION_ROOT).
        max_pages (int): Max pages to process from the start (for both modes).
        dpi (int): Render DPI used by OCR fallback (suggest 300–400).
        lang (str): Tesseract languages for OCR (e.g., 'eng+rus'); falls back to OCR_LANGS env or 'eng'.
        preprocess (str): Image preprocessing for OCR: 'auto'|'none'|'strong'.
        max_chars (int): Truncate the total returned text to this many characters.

    Returns:
        str: Extracted text (possibly truncated), '(no extractable text)', or 'error:<message>'.

    Notes:
        This tool first tries native text extraction via `read_pdf`. If that yields
        no text, it falls back to `read_pdf_ocr` (Tesseract + Poppler).
    """
    try:
        # 1) Try native text
        txt = read_pdf(path=path, max_pages=max_pages)  # returns "(no extractable text)" if scanned
        if isinstance(txt, str) and txt.strip() and "(no extractable text)" not in txt:
            return txt[:max_chars]

        # 2) Fallback to OCR
        return read_pdf_ocr(
            path=path, max_pages=max_pages, dpi=dpi,
            lang=lang, preprocess=preprocess, max_chars=max_chars
        )
    except Exception as e:
        return f"error:{e}"

def read_pdf(path: str, max_pages: int = 20) -> str:
    """Extract text from a non-scanned PDF (no OCR).

    Args:
        path (str): PDF path under action folder.
        max_pages (int): Max pages to extract.

    Returns:
        str: Text or '(no extractable text)' or 'error:<message>'.
    """
    try:
        from pypdf import PdfReader
        target = _normalize_to_root(path)
        reader = PdfReader(str(target))
        pages = []
        for i, page in enumerate(reader.pages):
            if i >= max_pages:
                break
            pages.append(page.extract_text() or "")
        txt = "\n\n".join(pages).strip()
        return txt if txt else "(no extractable text)"
    except Exception as e:
        return f"error:{e}"

def read_pdf_ocr(
    path: str,
    max_pages: int = 10,
    dpi: int = 300,
    lang: str = "",
    preprocess: str = "auto",
    max_chars: int = 20000,
) -> str:
    """OCR a scanned PDF by rendering pages to images then running Tesseract.

    Args:
        path (str): PDF path under action folder (relative to ACTION_ROOT).
        max_pages (int): Max pages from start; <=0 means ALL pages.
        dpi (int): Render DPI for OCR (suggest 300–400).
        lang (str): Tesseract languages (e.g., 'eng+rus'); falls back to OCR_LANGS env or 'eng'.
        preprocess (str): 'auto'|'none'|'strong' (strong adds binarization).
        max_chars (int): Truncate total output to this many characters.

    Returns:
        str: Extracted text (maybe truncated) or '(no text found)' or 'error:<message>'.
    """
    try:
        from pdf2image import convert_from_path

        # Resolve paths and env defaults
        target = _normalize_to_root(path)
        poppler_path = POPPLER_PATH or os.getenv("POPPLER_PATH") or None
        ocr_langs = (lang or os.getenv("OCR_LANGS", "") or DEFAULT_OCR_LANGS).strip() or "eng"

        # Tesseract binary (Windows)
        tcmd = os.getenv("TESSERACT_CMD")
        if tcmd:
            pytesseract.pytesseract.tesseract_cmd = tcmd

        # Clamp sane DPI
        dpi = max(100, min(int(dpi), 600))

        # pdf2image options
        kwargs = {"dpi": dpi}
        if poppler_path:
            kwargs["poppler_path"] = poppler_path

        # If max_pages <= 0 -> process all pages (no last_page bound)
        if isinstance(max_pages, int) and max_pages > 0:
            kwargs["first_page"] = 1
            kwargs["last_page"] = max_pages

        images = convert_from_path(str(target), **kwargs)

        out_parts = []
        total = 0

        for i, img in enumerate(images, start=1):
            page = img
            if preprocess != "none":
                page = ImageOps.grayscale(page)
                if preprocess in ("auto", "strong"):
                    page = ImageOps.autocontrast(page).filter(ImageFilter.MedianFilter(size=3))
                if preprocess == "strong":
                    # Simple thresholding for noisy scans
                    page = page.point(lambda p: 255 if p > 160 else 0)

            text = pytesseract.image_to_string(page, lang=ocr_langs)
            text = (text or "").strip()

            segment = f"[page {i}]\n{text}".strip()
            if segment:
                # Truncate progressively to respect max_chars
                remaining = max_chars - total
                if remaining <= 0:
                    break
                seg_trimmed = segment[: max(0, remaining)]
                out_parts.append(seg_trimmed)
                total += len(seg_trimmed)
                if total >= max_chars:
                    break

        final_txt = "\n\n".join(out_parts).strip()
        return final_txt if final_txt else "(no text found)"
    except Exception as e:
        return f"error:{e}"

@tool
def read_image_ocr(path: str, lang: str = "", preprocess: str = "auto", max_chars: int = 20000) -> str:
    """OCR any image using Tesseract.

    Args:
        path (str): Image path under action folder.
        lang (str): Tesseract languages (e.g., 'eng+rus').
        preprocess (str): 'auto'|'none'|'strong'.
        max_chars (int): Truncate output.

    Returns:
        str: Extracted text or '(no text found)' or 'error:<message>'.
    """
    try:
        target = _normalize_to_root(path)
        img = Image.open(str(target))
        if preprocess != "none":
            img = ImageOps.grayscale(img)
            if preprocess in ("auto", "strong"):
                img = ImageOps.autocontrast(img).filter(ImageFilter.MedianFilter(size=3))
            if preprocess == "strong":
                img = img.point(lambda p: 255 if p > 160 else 0)
        langs = (lang or DEFAULT_OCR_LANGS).strip() or "eng"
        text = pytesseract.image_to_string(img, lang=langs)
        return (text or "").strip()[:max_chars] or "(no text found)"
    except Exception as e:
        return f"error:{e}"

@tool
def read_docx(path: str, max_chars: int = 80000) -> str:
    """Read a .docx and return text with headings/lists preserved (Markdown-ish).

    Args:
        path (str): .docx path under the action folder (relative to ACTION_ROOT).
        max_chars (int): Truncate total output.

    Returns:
        str: Text content (maybe truncated) or '(no text found)' or 'error:<message>'.
    """
    try:
        from docx import Document
        target = _normalize_to_root(path)
        doc = Document(str(target))

        def _is_list(para):
            ppr = getattr(para._p, "pPr", None)
            return hasattr(ppr, "numPr") and ppr.numPr is not None if ppr is not None else False

        lines, total = [], 0
        for para in doc.paragraphs:
            txt = para.text.strip()
            if not txt:
                continue
            style = (para.style.name or "").lower()
            # Headings
            if "heading" in style:
                lvl = 1
                for n in ("1", "2", "3", "4", "5", "6"):
                    if n in style:
                        lvl = int(n); break
                line = f"{'#'*lvl} {txt}"
            # Lists
            elif _is_list(para) or "list" in style:
                line = f"- {txt}"
            else:
                line = txt

            remaining = max_chars - total
            if remaining <= 0:
                break
            line = line[: max(0, remaining)]
            lines.append(line)
            total += len(line) + 1

        # Simple table extraction (optional)
        for tbl in getattr(doc, "tables", []):
            row_texts = []
            for row in tbl.rows:
                cells = [" ".join(p.text for p in cell.paragraphs).strip() for cell in row.cells]
                row_texts.append("| " + " | ".join(cells) + " |")
            if row_texts:
                lines.append("\n".join(row_texts))

        out = "\n".join(lines).strip()
        return out if out else "(no text found)"
    except Exception as e:
        return f"error:{e}"

@tool
def read_xlsx(path: str, sheet: str = None, max_rows: int = 50) -> str:
    """Read a few rows from an XLSX sheet.

    Args:
        path (str): XLSX path under action folder.
        sheet (str): Sheet name to open (None = first sheet).
        max_rows (int): Max rows to return.

    Returns:
        str: Simple preview or '(empty sheet)' or 'error:<message>'.
    """
    try:
        from openpyxl import load_workbook
        target = _normalize_to_root(path)
        wb = load_workbook(str(target), read_only=True, data_only=True)
        ws = wb[sheet] if sheet and sheet in wb.sheetnames else wb[wb.sheetnames[0]]
        rows = []
        for i, row in enumerate(ws.iter_rows(values_only=True), 1):
            if i > max_rows:
                break
            rows.append(", ".join("" if v is None else str(v) for v in row))
        wb.close()
        return "\n".join(rows) if rows else "(empty sheet)"
    except Exception as e:
        return f"error:{e}"

@tool
def read_csv(path: str, delimiter: str = "", max_rows: int = 100) -> str:
    """Preview a CSV/TSV with delimiter detection.

    Args:
        path (str): CSV path under action folder.
        delimiter (str): Explicit delimiter to force (optional).
        max_rows (int): Max rows to return.

    Returns:
        str: Header + rows or '(empty csv)' or 'error:<message>'.
    """
    try:
        import csv
        target = _normalize_to_root(path)
        data = target.read_bytes()
        text = None
        for enc in ("utf-8", "cp1251", "utf-16", "latin-1"):
            try:
                text = data.decode(enc); break
            except Exception:
                continue
        if text is None:
            text = data.decode("utf-8", errors="replace")
        sniffer = csv.Sniffer()
        if delimiter:
            dialect = csv.excel; dialect.delimiter = delimiter  # type: ignore[attr-defined]
        else:
            try:
                first = next((ln for ln in text.splitlines() if ln.strip()), "")
                dialect = sniffer.sniff(first) if first else csv.excel
            except Exception:
                dialect = csv.excel
        out, reader = [], csv.reader(text.splitlines(), dialect=dialect)
        for i, row in enumerate(reader):
            if i > max_rows:
                break
            out.append(", ".join(row))
        return "\n".join(out) if out else "(empty csv)"
    except Exception as e:
        return f"error:{e}"

@tool
def count_files_by_type(
    path: str = ".",
    glob_pattern: str = "**/*",
    exclude: str = ".git,.venv,__pycache__,.ipynb_checkpoints,node_modules,dist,build"
) -> str:
    """Count files grouped by extension (deduped) with optional exclusions.

    Args:
        path (str): Subfolder under the action folder where counting begins.
        glob_pattern (str): Glob pattern to include (e.g., '**/*', '**/*.py').
        exclude (str): Comma-separated directory names to skip anywhere in paths
            (e.g., '.git,.venv,__pycache__,node_modules'). These are matched as
            path segments and excluded from the count.

    Returns:
        str: JSON dict mapping extension (e.g., '.py') → counts; '' means no extension.
    """
    try:
        import json
        cnt = Counter()
        base = _normalize_to_root(path)
        for p in _iter_files_unique(base, glob_pattern, exclude.split(",")):
            cnt[p.suffix or ""] += 1
        return json.dumps(
            dict(sorted(cnt.items(), key=lambda kv: (-kv[1], kv[0]))),
            ensure_ascii=False, indent=2
        )
    except Exception as e:
        return f"error:{e}"

@tool
def total_file_count(
    path: str = ".",
    glob_pattern: str = "**/*",
    exclude: str = ".git,.venv,__pycache__,.ipynb_checkpoints,node_modules,dist,build"
) -> str:
    """Count files recursively with optional exclusions.

    Args:
        path (str): Start under action folder.
        glob_pattern (str): Include pattern (e.g., '**/*' or '**/*.py').
        exclude (str): Comma-separated dir names to skip anywhere in the path.

    Returns:
        str: The number of files as text, or 'error:<msg>'.
    """
    try:
        base = _normalize_to_root(path)
        n = sum(1 for _ in _iter_files_unique(base, glob_pattern, exclude.split(",")))
        return str(n)
    except Exception as e:
        return f"error:{e}"

# --------- Tool registry ----------
def get_tools(safe_mode: bool, minimal: bool = True):
    """
    Return the list of tools exposed to the model.
    minimal=True keeps a single obvious tool per job to reduce confusion.
    """
    toolset = [
        cwd, list_dir, search_text, read_text, write_file,
        # Single choice for PDFs: auto-detect → fallback to OCR
        read_pdf_auto,
        # DOCX reader (Markdown-ish)
        read_docx,
        # Tabular/text helpers
        read_xlsx, read_csv,
        # OCR for images
        read_image_ocr,
        # Stats
        count_files_by_type,
        total_file_count,
    ]
    return toolset if safe_mode else toolset + [sh]
